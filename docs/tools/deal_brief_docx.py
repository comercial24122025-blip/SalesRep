#!/usr/bin/env python3
from __future__ import annotations

import json
import sys
from copy import deepcopy
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SKILL_ROOT = Path("/Users/erickmendez/.codex/plugins/cache/openai-primary-runtime/documents/26.430.10722/skills/documents")
sys.path.insert(0, str(SKILL_ROOT / "scripts"))

from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa  # type: ignore  # noqa: E402


KIND_CONFIG = {
    "legal": {
        "title": "New Service Agreement Request",
        "subtitle": "Legal intake and agreed commercials generated from SalesRep",
        "filename_suffix": "new-service-agreement-request",
    },
    "proposal": {
        "title": "Commercial Proposal",
        "subtitle": "Standard Evolution commercial proposal prepared from Spazio",
        "filename_suffix": "commercial-proposal",
    },
    "dd": {
        "title": "DD Request",
        "subtitle": "Due diligence intake generated from SalesRep",
        "filename_suffix": "dd-request",
    },
    "integration": {
        "title": "Integration Request",
        "subtitle": "Operational integration brief generated from SalesRep",
        "filename_suffix": "integration-request",
    },
    "signoff": {
        "title": "Legal Signoff Request",
        "subtitle": "Final internal signoff generated from SalesRep",
        "filename_suffix": "legal-signoff-request",
    },
}

EVOLUTION_GROUP_PARAGRAPHS = [
    "The Evolution Group is a leading provider of online gaming content, bringing together premium live casino, RNG, slots, and branded content under one commercial relationship.",
    "Evolution offers the widest range of Live Casino games with one-of-a-kind innovations, immersive classic tables, award-winning game shows, and a multi-studio footprint that supports regulated and growth markets across LATAM and beyond.",
    "To complement the core live casino proposition, the Group also provides access to premium RNG and slot content from leading studios, creating a broader casino solution for operators looking to scale player acquisition, retention, and wallet share.",
]

EVOLUTION_BRAND_SECTIONS = [
    (
        "Evolution",
        [
            "Evolution provides the core live casino proposition, including roulette, blackjack, baccarat, game shows, and native-table options aligned to market preferences and localization needs.",
            "For LATAM opportunities, Evolution can also support native environments, premium tables, and dedicated positioning opportunities depending on the agreed commercial structure.",
        ],
    ),
    (
        "Ezugi",
        [
            "Ezugi complements the live casino offering with additional live tables and regional depth across multiple jurisdictions, helping operators broaden table variety and player coverage.",
        ],
    ),
    (
        "NetEnt and Red Tiger",
        [
            "NetEnt and Red Tiger add premium slot and branded content to the commercial package, with proven top-performing titles, regular launches, and access to enhanced promotional mechanics where applicable.",
        ],
    ),
    (
        "Big Time Gaming and Nolimit City",
        [
            "Big Time Gaming and Nolimit City extend the content stack with high-recognition mechanics, distinctive math models, and high-engagement slots that can strengthen the broader casino proposition.",
        ],
    ),
]


def clean(value) -> str:
    return " ".join(str(value or "").replace("\u00a0", " ").split())


def clean_multiline(value) -> str:
    return "\n".join(
        line for line in (" ".join(str(part).split()) for part in str(value or "").replace("\u00a0", " ").splitlines()) if line
    )


def first_value(*values) -> str:
    for value in values:
        text = clean(value)
        if text:
            return text
    return "TBC"


def document_client_name(deal: dict) -> str:
    return first_value(
        deal.get("documentClientName"),
        deal.get("companyName"),
        deal.get("client"),
        deal.get("operator"),
        deal.get("deal"),
    )


def to_positive_int(value, default: int) -> int:
    try:
        number = int(str(value or "").strip())
        return number if number > 0 else default
    except Exception:
        return default


def parse_iso_date(value: str | None) -> date | None:
    text = clean(value)
    if not text:
        return None
    try:
        return date.fromisoformat(text)
    except ValueError:
        return None


def proposal_validity_days(deal: dict) -> int:
    return to_positive_int(deal.get("proposalValidityDays"), 30)


def proposal_issue_date(deal: dict) -> date:
    return parse_iso_date(deal.get("offerDate")) or date.today()


def proposal_valid_until(deal: dict) -> date:
    return parse_iso_date(deal.get("proposalValidUntil")) or (proposal_issue_date(deal) + timedelta(days=proposal_validity_days(deal)))


def proposal_validity_text(deal: dict) -> str:
    days = proposal_validity_days(deal)
    month_hint = " (1 month)" if 28 <= days <= 31 else ""
    return f"{days} days{month_hint}, valid until {proposal_valid_until(deal).strftime('%B %d, %Y')}"


def parse_commercial_schedule(value) -> list[dict[str, str]]:
    rows = []
    for line in clean_multiline(value).splitlines():
        parts = [clean(part) for part in line.split("|")]
        if len(parts) >= 4:
            rows.append(
                {
                    "product": parts[0],
                    "tier": parts[1],
                    "commercial": parts[2],
                    "notes": " | ".join(parts[3:]),
                }
            )
        elif len(parts) == 3:
            rows.append(
                {
                    "product": parts[0],
                    "tier": parts[1],
                    "commercial": parts[2],
                    "notes": "",
                }
            )
        elif len(parts) == 2:
            rows.append(
                {
                    "product": parts[0],
                    "tier": "",
                    "commercial": parts[1],
                    "notes": "",
                }
            )
        elif len(parts) == 1 and parts[0]:
            rows.append(
                {
                    "product": parts[0],
                    "tier": "",
                    "commercial": "",
                    "notes": "",
                }
            )
    return rows


def load_payload(path: Path) -> dict:
    raw = json.loads(path.read_text("utf-8"))
    if isinstance(raw, dict) and isinstance(raw.get("deal"), dict):
      return raw["deal"]
    return raw if isinstance(raw, dict) else {}


def clear_body_preserve_layout(doc: Document) -> None:
    body = doc._body._element
    preserved_children = []
    for para in doc.paragraphs[:6]:
        if "w:drawing" in para._element.xml:
            preserved_children.append(deepcopy(para._element))
    sect_pr = body.sectPr
    for child in list(body):
        body.remove(child)
    for child in preserved_children[:1]:
        body.append(child)
    if sect_pr is not None:
        body.append(sect_pr)


def set_document_meta(doc: Document, title: str, client_name: str) -> None:
    doc.core_properties.title = title
    doc.core_properties.subject = client_name
    doc.core_properties.category = "Spazio Brief"
    doc.core_properties.comments = "Generated from Spazio"


def add_title_block(doc: Document, kind: str, deal: dict) -> None:
    config = KIND_CONFIG[kind]
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run(config["title"])
    title_run.bold = True
    title_run.font.size = Pt(25)
    title_run.font.color.rgb = RGBColor(0x16, 0x1A, 0x1F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle_run = subtitle.add_run(
        f"{document_client_name(deal)} · "
        f"{first_value(deal.get('market'))} · "
        f"{date.today().strftime('%B %d, %Y')}"
    )
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(15)
    subtitle_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    intro = doc.add_paragraph(style="Normal")
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_before = Pt(0)
    intro.paragraph_format.space_after = Pt(10)
    intro_run = intro.add_run(config["subtitle"])
    intro_run.italic = True
    intro_run.font.size = Pt(11)


def add_metadata_table(doc: Document, deal: dict) -> None:
    rows = [
        (
            ("Deal", first_value(deal.get("deal"))),
            ("Client", document_client_name(deal)),
        ),
        (
            ("Operator", first_value(deal.get("operator"))),
            ("Market", first_value(deal.get("market"))),
        ),
        (
            ("Platform", first_value(deal.get("platform"))),
            ("Stage", first_value(deal.get("stage"))),
        ),
        (
            ("KAM", first_value(deal.get("kam"))),
            ("Deal Value (EUR)", first_value(deal.get("dealValue"))),
        ),
    ]

    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    widths = column_widths_from_weights([0.14, 0.36, 0.16, 0.34], section_content_width_dxa(doc.sections[0]))
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)

    for row, pair_group in zip(table.rows, rows):
        flattened = [item for pair in pair_group for item in pair]
        for cell, value in zip(row.cells, flattened):
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles["Normal"]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
        for label_cell in (row.cells[0], row.cells[2]):
            for paragraph in label_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    style = f"Heading {level}"
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)


def add_kv_paragraph(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(value)


def add_section_table(doc: Document, pairs: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    widths = column_widths_from_weights([0.3, 0.7], section_content_width_dxa(doc.sections[0]))
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)
    for row, (label, value) in zip(table.rows, pairs):
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in row.cells[0].paragraphs:
            if paragraph.runs:
                paragraph.runs[0].bold = True


def add_text_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.add_run(text)


def add_commercial_schedule_table(doc: Document, schedule_rows: list[dict[str, str]]) -> None:
    if not schedule_rows:
        return

    add_heading(doc, "Commercial Schedule", 2)
    table = doc.add_table(rows=len(schedule_rows) + 1, cols=4)
    table.style = "Table Grid"
    widths = column_widths_from_weights([0.24, 0.18, 0.34, 0.24], section_content_width_dxa(doc.sections[0]))
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)

    headers = ["Product / Scope", "Tier / Volume", "Commercial", "Notes"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row, item in zip(table.rows[1:], schedule_rows):
        values = [item["product"], item["tier"], item["commercial"], item["notes"]]
        for cell, value in zip(row.cells, values):
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles["Normal"]


def add_legal_content(doc: Document, deal: dict) -> None:
    schedule_rows = parse_commercial_schedule(deal.get("commercialSchedule"))
    add_heading(doc, "New Service Agreement Summary", 1)
    add_text_block(
        doc,
        first_value(
            deal.get("statusText"),
            deal.get("comments"),
            "Please start legal review and contract preparation for this client opportunity.",
        ),
    )

    add_heading(doc, "Company and Registration Details", 2)
    add_section_table(
        doc,
        [
            ("Company Name", first_value(deal.get("companyName"), deal.get("client"), deal.get("operator"))),
            ("Legal Entity", first_value(deal.get("legalEntity"))),
            ("Registration Number", first_value(deal.get("companyRegistrationNumber"))),
            ("Registered Address", first_value(deal.get("companyRegisteredAddress"))),
            ("License", first_value(deal.get("companyLicense"), deal.get("licenseStatus"))),
            ("Company Legal Representative", first_value(deal.get("companyLegalRepresentative"), deal.get("legalRepresentativeName"))),
        ],
    )

    add_heading(doc, "Billing and Operational Contacts", 2)
    add_section_table(
        doc,
        [
            ("Invoice Email", first_value(deal.get("invoiceEmail"))),
            ("Support Email", first_value(deal.get("supportEmail"))),
            ("Management Email", first_value(deal.get("managementEmail"))),
            ("Primary Contact", first_value(deal.get("primaryContact"))),
            ("Decision Maker", first_value(deal.get("decisionMaker"))),
            ("DD Contact", first_value(f"{clean(deal.get('ddContactName'))} {clean(deal.get('ddContactEmail'))}")),
        ],
    )

    add_heading(doc, "Legal Representative", 2)
    add_section_table(
        doc,
        [
            ("Full Name", first_value(deal.get("legalRepresentativeName"), deal.get("companyLegalRepresentative"))),
            ("ID", first_value(deal.get("legalRepresentativeId"))),
            ("Address", first_value(deal.get("legalRepresentativeAddress"))),
            ("Email", first_value(deal.get("legalRepresentativeEmail"))),
        ],
    )

    add_heading(doc, "Commercial Context", 2)
    add_kv_paragraph(doc, "Proposal Request", first_value(deal.get("proposalRequest")))
    add_kv_paragraph(doc, "Products in Scope", first_value(deal.get("negotiatedProducts"), deal.get("productsPotential"), deal.get("productsCurrent")))
    add_kv_paragraph(doc, "Commercial Terms", first_value(deal.get("commercialTerms")))
    add_kv_paragraph(doc, "Pricing Base", first_value(deal.get("pricingBase")))
    add_kv_paragraph(doc, "Deduction Terms", first_value(deal.get("deductionTerms")))
    add_kv_paragraph(doc, "Setup Fee Status", first_value(deal.get("setupFeeStatus")))
    add_kv_paragraph(doc, "Setup Fee Amount (EUR)", first_value(deal.get("setupFeeAmount")))
    add_kv_paragraph(doc, "Marketing Commitments", first_value(deal.get("marketingCommitments")))
    add_kv_paragraph(doc, "Live Games Top Position", first_value(deal.get("liveGamesTopPosition")))
    add_kv_paragraph(doc, "Slots Top Position", first_value(deal.get("slotsTopPosition")))
    add_kv_paragraph(doc, "Action Items", first_value(deal.get("actionItems")))

    add_heading(doc, "Deductions Allowed", 2)
    add_section_table(
        doc,
        [
            ("Deductions Allowed", first_value(deal.get("deductionsAllowed"))),
            ("Bonus Cap", first_value(deal.get("bonusCap"))),
            ("Gaming Tax", first_value(deal.get("gamingTax"))),
            ("Withholding", first_value(deal.get("withholdingTax"))),
            ("Advance Payment", first_value(deal.get("advancePayment"))),
            ("Credit Notes", first_value(deal.get("creditNotes"))),
        ],
    )

    add_commercial_schedule_table(doc, schedule_rows)


def add_proposal_content(doc: Document, deal: dict) -> None:
    schedule_rows = parse_commercial_schedule(deal.get("commercialSchedule"))
    add_heading(doc, "Executive Summary", 1)
    add_text_block(
        doc,
        first_value(
            deal.get("proposalRequest"),
            deal.get("statusText"),
            f"This commercial proposal outlines the Evolution standard offering, commercials, and implementation scope for {document_client_name(deal)}.",
        ),
    )

    add_heading(doc, "About the Evolution Group", 1)
    for paragraph in EVOLUTION_GROUP_PARAGRAPHS:
        add_text_block(doc, paragraph)

    add_heading(doc, "About our brands", 1)
    for title, paragraphs in EVOLUTION_BRAND_SECTIONS:
        add_heading(doc, title, 2)
        for paragraph in paragraphs:
            add_text_block(doc, paragraph)

    add_heading(doc, "Our Commercial Proposal", 1)
    add_text_block(
        doc,
        first_value(
            deal.get("negotiationScope"),
            "The commercial proposal below reflects the requested scope, pricing structure, positioning commitments, and operational assumptions currently in discussion.",
        ),
    )

    add_heading(doc, "Client Opportunity", 2)
    add_kv_paragraph(doc, "Segment", first_value(deal.get("segment"), deal.get("type")))
    add_kv_paragraph(doc, "Strategic Fit", first_value(deal.get("strategicFit")))
    add_kv_paragraph(doc, "Revenue Potential EUR", first_value(deal.get("revenuePotentialEur"), deal.get("dealValue")))
    add_kv_paragraph(doc, "Website", first_value(deal.get("url"), deal.get("siteStatus")))
    add_kv_paragraph(doc, "Proposal Validity", proposal_validity_text(deal))
    add_kv_paragraph(doc, "Market", first_value(deal.get("market")))
    add_kv_paragraph(doc, "Commercial Owner", first_value(deal.get("kam")))

    add_heading(doc, "Our Offering and Pricing", 1)
    add_heading(doc, "Requested Offering", 2)
    add_kv_paragraph(doc, "Requested Products", first_value(deal.get("negotiatedProducts"), deal.get("productsPotential"), deal.get("productsCurrent")))
    add_kv_paragraph(doc, "Other Live Suppliers", first_value(deal.get("otherLiveSuppliers")))
    add_kv_paragraph(doc, "Activation Requirements", first_value(deal.get("activationRequirements")))
    add_kv_paragraph(doc, "Marketing Commitments", first_value(deal.get("marketingCommitments")))
    add_kv_paragraph(doc, "Live Games Positioning", first_value(deal.get("liveGamesTopPosition")))
    add_kv_paragraph(doc, "Slots Positioning", first_value(deal.get("slotsTopPosition")))

    add_heading(doc, "Commercial Terms and Pricing", 2)
    add_section_table(
        doc,
        [
            ("Commercial Terms", first_value(deal.get("commercialTerms"))),
            ("Pricing Base", first_value(deal.get("pricingBase"))),
            ("Deduction Terms", first_value(deal.get("deductionTerms"))),
            ("Negotiation Scope", first_value(deal.get("negotiationScope"))),
            ("Setup Fee Status", first_value(deal.get("setupFeeStatus"))),
            ("Setup Fee Amount", first_value(deal.get("setupFeeAmount"))),
        ],
    )
    add_commercial_schedule_table(doc, schedule_rows)

    add_heading(doc, "Deductions and Financial Conditions", 2)
    add_section_table(
        doc,
        [
            ("Deductions Allowed", first_value(deal.get("deductionsAllowed"))),
            ("Bonus Cap", first_value(deal.get("bonusCap"))),
            ("Gaming Tax", first_value(deal.get("gamingTax"))),
            ("Withholding", first_value(deal.get("withholdingTax"))),
            ("Advance Payment", first_value(deal.get("advancePayment"))),
            ("Credit Notes", first_value(deal.get("creditNotes"))),
        ],
    )

    add_heading(doc, "Implementation Dependencies", 2)
    add_kv_paragraph(doc, "Integration Request", first_value(deal.get("integrationRequest")))
    add_kv_paragraph(doc, "Integration Team", first_value(deal.get("integrationTeam")))
    add_kv_paragraph(doc, "Jira", first_value(deal.get("jira")))
    add_kv_paragraph(doc, "Next Actions", first_value(deal.get("actionItems"), deal.get("updates")))

    add_heading(doc, "Appendices Available on Request", 2)
    add_text_block(doc, "Appendix 1: Evolution Generic Games Offering")
    add_text_block(doc, "Appendix 2: NetEnt and Red Tiger branded and premium slots")
    add_text_block(doc, "Appendix 3: Ezugi generic and premium content")
    add_text_block(doc, "Appendix 4: LATAM package")

    add_heading(doc, "Closing", 2)
    add_text_block(
        doc,
        "We trust this proposal aligns with the opportunity in scope and provides a strong commercial foundation to move the discussion forward.",
    )
    add_text_block(
        doc,
        "We remain committed to supporting the implementation process with the required legal, technical, and commercial coordination to help deliver a successful launch.",
    )
    add_text_block(doc, "Kind regards,")
    add_text_block(doc, first_value(deal.get("kam"), "Evolution LATAM Commercial Team"))
    add_text_block(doc, "Evolution LATAM")


def add_dd_content(doc: Document, deal: dict) -> None:
    add_heading(doc, "Due Diligence Request Summary", 1)
    add_text_block(
        doc,
        first_value(
            deal.get("comments"),
            deal.get("statusText"),
            "Please initiate due diligence review for this client opportunity.",
        ),
    )

    add_heading(doc, "Client and Corporate Information", 2)
    add_section_table(
        doc,
        [
            ("Client Name", document_client_name(deal)),
            ("Legal Entity", first_value(deal.get("legalEntity"))),
            ("Registration Number", first_value(deal.get("companyRegistrationNumber"))),
            ("Registered Address", first_value(deal.get("companyRegisteredAddress"))),
            ("License", first_value(deal.get("companyLicense"), deal.get("licenseStatus"))),
            ("DD Start Date", first_value(deal.get("ddDate"))),
        ],
    )

    add_heading(doc, "DD Contacts", 2)
    add_section_table(
        doc,
        [
            ("DD Contact Name", first_value(deal.get("ddContactName"))),
            ("DD Contact Email", first_value(deal.get("ddContactEmail"))),
            ("Legal Representative", first_value(deal.get("legalRepresentativeName"), deal.get("companyLegalRepresentative"))),
            ("Legal Representative Email", first_value(deal.get("legalRepresentativeEmail"))),
            ("Primary Contact", first_value(deal.get("primaryContact"))),
            ("Decision Maker", first_value(deal.get("decisionMaker"))),
        ],
    )

    add_heading(doc, "Commercial Context", 2)
    add_kv_paragraph(doc, "Products in Scope", first_value(deal.get("negotiatedProducts"), deal.get("productsPotential"), deal.get("productsCurrent")))
    add_kv_paragraph(doc, "Commercial Terms", first_value(deal.get("commercialTerms")))
    add_kv_paragraph(doc, "Deductions Allowed", first_value(deal.get("deductionsAllowed"), deal.get("deductionTerms")))
    add_kv_paragraph(doc, "Action Items", first_value(deal.get("actionItems")))
    add_kv_paragraph(doc, "Notes", first_value(deal.get("statusText"), deal.get("comments")))


def add_integration_content(doc: Document, deal: dict) -> None:
    add_heading(doc, "Integration Request Summary", 1)
    add_text_block(
        doc,
        first_value(
            deal.get("integrationRequest"),
            "Please start operational support and technical scoping for this client integration.",
        ),
    )

    add_heading(doc, "Client Background", 2)
    add_section_table(
        doc,
        [
            ("Client Type", first_value(deal.get("type"))),
            ("Client Based", first_value(deal.get("clientBased"), deal.get("jurisdiction"))),
            ("Legal Entity", first_value(deal.get("legalEntity"))),
            ("Market", first_value(deal.get("market"))),
            ("Platform", first_value(deal.get("platform"))),
            ("KAM", first_value(deal.get("kam"))),
        ],
    )

    add_heading(doc, "Technical and Commercial Scope", 2)
    add_kv_paragraph(doc, "Products", first_value(deal.get("negotiatedProducts"), deal.get("productsPotential"), deal.get("productsCurrent")))
    add_kv_paragraph(doc, "URL", first_value(deal.get("url"), deal.get("siteStatus")))
    add_kv_paragraph(doc, "Other Live Suppliers", first_value(deal.get("otherLiveSuppliers")))
    add_kv_paragraph(doc, "Integration Request", first_value(deal.get("integrationRequest")))
    add_kv_paragraph(doc, "Commercial Terms", first_value(deal.get("commercialTerms")))

    add_heading(doc, "Key Contacts and Channels", 2)
    add_section_table(
        doc,
        [
            ("Integration Team", first_value(deal.get("integrationTeam"))),
            ("Integration Email", first_value(deal.get("integrationEmail"))),
            ("Integration Chat", first_value(deal.get("skype"))),
            ("Teams Group", first_value(deal.get("teamsGroup"))),
            ("Jira Ticket", first_value(deal.get("jira"))),
            ("DD Ticket", first_value(deal.get("ddTicket"))),
        ],
    )

    add_heading(doc, "Launch Dependencies", 2)
    add_kv_paragraph(doc, "License", first_value(deal.get("companyLicense"), deal.get("licenseStatus")))
    add_kv_paragraph(doc, "Action Items", first_value(deal.get("actionItems")))
    add_kv_paragraph(doc, "Other Info", first_value(deal.get("otherInfo"), deal.get("entityInfo"), deal.get("updates")))


def add_signoff_content(doc: Document, deal: dict) -> None:
    add_heading(doc, "Legal Signoff Summary", 1)
    add_text_block(
        doc,
        first_value(
            deal.get("legalSignoffRequest"),
            deal.get("statusText"),
            deal.get("comments"),
            "Please complete final legal signoff before authorizing Go Live.",
        ),
    )

    add_heading(doc, "Approval Readiness", 2)
    add_section_table(
        doc,
        [
            ("Client Name", document_client_name(deal)),
            ("Legal Approval Date", first_value(deal.get("legalApprovalDate"))),
            ("Legal Status", first_value(deal.get("legalStatus"))),
            ("DD Status", first_value(deal.get("ddStatus"))),
            ("Integration Status", first_value(deal.get("integrationStatus"))),
            ("Go Live Status", first_value(deal.get("goLiveStatus"))),
        ],
    )

    add_heading(doc, "Commercial and Launch Scope", 2)
    add_kv_paragraph(doc, "Products in Scope", first_value(deal.get("negotiatedProducts"), deal.get("productsPotential"), deal.get("productsCurrent")))
    add_kv_paragraph(doc, "Commercial Terms", first_value(deal.get("commercialTerms")))
    add_kv_paragraph(doc, "Proposal Request", first_value(deal.get("proposalRequest")))
    add_kv_paragraph(doc, "Integration Request", first_value(deal.get("integrationRequest")))
    add_kv_paragraph(doc, "Live Date", first_value(deal.get("liveDate"), deal.get("liveSince")))

    add_heading(doc, "Execution Trace", 2)
    add_section_table(
        doc,
        [
            ("Jira Ticket", first_value(deal.get("jira"))),
            ("DD Ticket", first_value(deal.get("ddTicket"))),
            ("Action Items", first_value(deal.get("actionItems"))),
            ("Updates", first_value(deal.get("updates"))),
            ("Other Info", first_value(deal.get("otherInfo"), deal.get("entityInfo"))),
        ],
    )


def build_doc(kind: str, deal: dict, template_path: Path, output_path: Path) -> None:
    doc = Document(str(template_path))
    clear_body_preserve_layout(doc)
    set_document_meta(doc, KIND_CONFIG[kind]["title"], document_client_name(deal))
    add_title_block(doc, kind, deal)
    add_metadata_table(doc, deal)

    if kind == "legal":
        add_legal_content(doc, deal)
    elif kind == "proposal":
        add_proposal_content(doc, deal)
    elif kind == "dd":
        add_dd_content(doc, deal)
    elif kind == "integration":
        add_integration_content(doc, deal)
    elif kind == "signoff":
        add_signoff_content(doc, deal)
    else:
        raise ValueError(f"Unsupported kind: {kind}")

    doc.save(str(output_path))


def main() -> int:
    if len(sys.argv) != 5:
        print("Usage: deal_brief_docx.py <kind> <payload_json> <output_docx> <template_docx>", file=sys.stderr)
        return 1

    kind = clean(sys.argv[1]).lower()
    payload_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])
    template_path = Path(sys.argv[4])

    if kind not in KIND_CONFIG:
        print(f"Unsupported kind: {kind}", file=sys.stderr)
        return 1
    if not payload_path.exists():
        print(f"Payload file not found: {payload_path}", file=sys.stderr)
        return 1
    if not template_path.exists():
        print(f"Template file not found: {template_path}", file=sys.stderr)
        return 1

    deal = load_payload(payload_path)
    build_doc(kind, deal, template_path, output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
